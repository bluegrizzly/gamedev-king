import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple

from docx import Document

from pdf_export import get_doc_output_dir

MAX_FILENAME_LEN = 120
ALLOWED_FILENAME_RE = re.compile(r"[^a-zA-Z0-9\-_. ]+")


def sanitize_docx_filename(value: str) -> str:
    cleaned = value.replace("\\", "_").replace("/", "_").strip()
    cleaned = cleaned.replace("..", "_")
    cleaned = ALLOWED_FILENAME_RE.sub("_", cleaned)
    cleaned = cleaned.strip(" .")
    if not cleaned:
        cleaned = "document"
    if not cleaned.lower().endswith(".docx"):
        cleaned = f"{cleaned}.docx"
    if len(cleaned) > MAX_FILENAME_LEN:
        base = cleaned[: MAX_FILENAME_LEN - 5].rstrip(" ._")
        cleaned = f"{base}.docx"
    return cleaned


def build_docx_filename(title: str) -> str:
    base = title.strip() or "document"
    base = base.lower().replace(" ", "_")
    base = ALLOWED_FILENAME_RE.sub("_", base)
    base = re.sub(r"_+", "_", base).strip("_")
    if not base:
        base = "document"
    timestamp = datetime.utcnow().strftime("%Y-%m-%d_%H%M%S")
    return sanitize_docx_filename(f"{base}_{timestamp}.docx")


def ensure_docx_output_path(filename: str, project_key: Optional[str] = None) -> Path:
    output_dir = get_doc_output_dir(project_key).resolve()
    candidate = (output_dir / filename).resolve()
    if output_dir not in candidate.parents and candidate != output_dir:
        raise ValueError("Invalid filename path.")
    return candidate


def _add_paragraph(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        document.add_paragraph("")
        return
    if stripped.startswith("### "):
        document.add_heading(stripped[4:].strip(), level=3)
        return
    if stripped.startswith("## "):
        document.add_heading(stripped[3:].strip(), level=2)
        return
    if stripped.startswith("# "):
        document.add_heading(stripped[2:].strip(), level=1)
        return
    if stripped.startswith(("- ", "* ")):
        document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return
    if re.match(r"^\d+\.\s+", stripped):
        document.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        return
    document.add_paragraph(stripped)


def write_docx(
    title: str,
    content: str,
    filename: str,
    project_key: Optional[str] = None,
) -> Tuple[str, Path]:
    safe_name = sanitize_docx_filename(filename)
    output_path = ensure_docx_output_path(safe_name, project_key)

    document = Document()
    document.add_heading(title, level=1)
    for line in content.splitlines():
        _add_paragraph(document, line)
    document.save(str(output_path))

    return safe_name, output_path


def run_export_docx_tool(args: dict) -> dict:
    title = str(args.get("title", "")).strip()
    content = str(args.get("content", "")).strip()
    filename = str(args.get("filename", "")).strip() if args.get("filename") else ""
    project_key = str(args.get("project_key", "")).strip() if args.get("project_key") else None

    if not title:
        raise ValueError("Tool arg 'title' is required.")
    if not content:
        raise ValueError("Tool arg 'content' is required.")
    if len(content.encode("utf-8")) > 2 * 1024 * 1024:
        raise ValueError("Tool arg 'content' exceeds 2MB limit.")

    final_filename = filename or build_docx_filename(title)
    saved_name, output_path = write_docx(title, content, final_filename, project_key)
    download_url = f"/downloads/{saved_name}"
    if project_key:
        download_url = f"{download_url}?project_key={project_key}"
    return {
        "ok": True,
        "filename": saved_name,
        "path": str(output_path),
        "download_url": download_url,
    }
